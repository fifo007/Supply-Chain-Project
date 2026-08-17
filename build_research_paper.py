from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parent
DATA = ROOT / "APL_Logistics.csv"
OUTPUT = ROOT / "APL_Logistics_Research_Paper.docx"
FIGURES = ROOT / "paper_figures"
PROFIT = "Order Profit Per Order"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "595959"


def set_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side in ("top", "start", "bottom", "end"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "80" if side in ("top", "bottom") else "120")
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)


def add_text(doc, text, bold_label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.33
    if bold_label and text.startswith(bold_label):
        label = p.add_run(bold_label)
        set_font(label, bold=True)
        rest = p.add_run(text[len(bold_label):])
        set_font(rest)
    else:
        set_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 6)
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(text), size=9, italic=True, color=GRAY)


def money(value):
    return f"${value / 1_000_000:.2f}M"


def percent(value):
    return f"{value:.1f}%"


def create_figures(df, market, segments, discount):
    FIGURES.mkdir(exist_ok=True)
    font = ImageFont.truetype("arial.ttf", 25)
    title_font = ImageFont.truetype("arialbd.ttf", 31)

    def bar_chart(labels, values, title, unit, color, path, vertical=False):
        image = Image.new("RGB", (1260, 620), "white")
        draw = ImageDraw.Draw(image)
        draw.text((45, 30), title, fill="#0B2545", font=title_font)
        max_value = max(values) * 1.12
        if vertical:
            left, right, top, bottom = 105, 1200, 130, 520
            width = (right - left) / len(labels)
            for i, (label, value) in enumerate(zip(labels, values)):
                x1, x2 = left + i * width + 35, left + (i + 1) * width - 35
                y1 = bottom - (value / max_value) * (bottom - top)
                draw.rectangle((x1, y1, x2, bottom), fill=color)
                draw.text((x1, y1 - 35), f"{value:.1f}{unit}", fill="#333333", font=font)
                draw.text((x1, bottom + 18), label, fill="#333333", font=font)
            draw.line((left, bottom, right, bottom), fill="#777777", width=2)
        else:
            left, right, top, row = 260, 1160, 145, 80
            for i, (label, value) in enumerate(zip(labels, values)):
                y = top + i * row
                draw.text((45, y + 9), label, fill="#333333", font=font)
                x2 = left + (value / max_value) * (right - left)
                draw.rounded_rectangle((left, y, x2, y + 42), radius=8, fill=color)
                draw.text((x2 + 12, y + 7), f"{value:.2f}{unit}", fill="#333333", font=font)
        image.save(path)

    ordered = market.sort_values(PROFIT)
    market_path = FIGURES / "market_profit.png"
    bar_chart(list(ordered.index), list(ordered[PROFIT] / 1_000_000), "Profit by market", "M", "#2E74B5", market_path)
    ordered = segments.sort_values(PROFIT)
    segment_path = FIGURES / "segment_profit.png"
    bar_chart(list(ordered.index), list(ordered[PROFIT] / 1_000_000), "Profit by customer segment", "M", "#9B1C1C", segment_path, vertical=True)
    ordered = discount.sort_values("Discount band")
    discount_path = FIGURES / "discount_margin.png"
    bar_chart([str(x) for x in ordered["Discount band"]], list(ordered["Margin %"]), "Profit margin by discount band", "%", "#1F4D78", discount_path, vertical=True)
    return market_path, segment_path, discount_path


def main():
    df = pd.read_csv(DATA, encoding="latin-1", low_memory=False)
    df["Discount %"] = pd.to_numeric(df["Order Item Discount Rate"], errors="coerce") * 100
    df["Sales"] = pd.to_numeric(df["Sales"], errors="coerce")
    df[PROFIT] = pd.to_numeric(df[PROFIT], errors="coerce")

    revenue = df["Sales"].sum()
    profit = df[PROFIT].sum()
    margin = profit / revenue * 100
    market = df.groupby("Market")[["Sales", PROFIT]].sum()
    segments = df.groupby("Customer Segment")[["Sales", PROFIT]].sum()
    categories = df.groupby("Category Name").agg({"Sales": "sum", PROFIT: "sum"}).sort_values(PROFIT, ascending=False)
    bands = pd.cut(df["Discount %"], [-0.01, 0, 10, 20, 100], labels=["No discount", "1-10%", "11-20%", ">20%"])
    discount = df.assign(**{"Discount band": bands}).groupby("Discount band", observed=False).agg({"Sales": "sum", PROFIT: "sum"}).reset_index()
    discount["Margin %"] = discount[PROFIT] / discount["Sales"] * 100
    market_path, segment_path, discount_path = create_figures(df, market, segments, discount)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.33

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("APL Logistics Profitability Intelligence | Research Paper"), size=8, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Supply Chain Analytics Project"), size=8, color=GRAY)

    # Editorial cover
    doc.add_paragraph().paragraph_format.space_after = Pt(78)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("RESEARCH PAPER"), size=11, bold=True, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("APL Logistics Profitability Intelligence"), size=28, bold=True, color="0B2545")
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Exploratory Data Analysis, Operational Insights, and Recommendations"), size=14, color=GRAY)
    p.paragraph_format.space_after = Pt(72)
    for line in ["Prepared by: [Student Name]", "Project: Supply Chain Analytics Dashboard", "Date: August 2026"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(line), size=11, color=GRAY)
        p.paragraph_format.space_after = Pt(5)

    doc.add_page_break()
    add_heading(doc, "Abstract")
    add_text(doc, f"This applied analytics study examines {len(df):,} APL Logistics order-line records to identify the commercial factors associated with revenue, profitability, customer value, product performance, and discounting. The dataset contains {len(df.columns)} fields spanning {df['Market'].nunique()} markets, {df['Category Name'].nunique()} product categories, and {df['Customer Id'].nunique():,} customers. Total recorded sales were {money(revenue)} and total profit was {money(profit)}, producing an overall profit margin of {percent(margin)}. The analysis was implemented as a Streamlit dashboard using Python, Pandas, and Plotly so that decision makers can interrogate the same indicators interactively. Results show strong market concentration in Europe and LATAM, variation in profit contribution across customer segments, and declining margin as discount intensity increases. The paper translates these findings into practical recommendations for pricing governance, customer strategy, and performance monitoring.")

    add_heading(doc, "1. Introduction")
    add_text(doc, "Supply-chain decision making depends on connecting commercial activity with financial outcomes. A dashboard can make these relationships visible by bringing sales, profit, customer, product, discount, and market indicators into a single analytical view. This project develops an interactive profitability-intelligence dashboard for APL Logistics and uses exploratory data analysis (EDA) to frame actionable management questions.")
    add_text(doc, "The purpose of the study is not to establish causality. Instead, it provides a descriptive baseline that helps stakeholders prioritize investigation and commercial improvement. The core questions are: Which customers, markets, products, and categories generate the most value? How are discounts associated with margin? What actions should managers take to improve profitability while protecting customer value?")

    add_heading(doc, "2. Data and Methodology")
    add_text(doc, "The analysis uses the supplied APL_Logistics.csv file. Each record is treated as an order-line item. The file was read using Latin-1 encoding, numeric fields were standardized, and derived measures were calculated for discount percentage, profit margin, and customer name. Records were aggregated by market, category, customer segment, product, and discount band.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = ["Metric", "Value"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "F2F4F7")
        set_font(cell.paragraphs[0].add_run(value), bold=True, color=DARK_BLUE)
    rows = [
        ("Records analysed", f"{len(df):,}"),
        ("Variables", f"{len(df.columns)}"),
        ("Customers", f"{df['Customer Id'].nunique():,}"),
        ("Markets", f"{df['Market'].nunique()}"),
        ("Product categories", f"{df['Category Name'].nunique()}"),
        ("Total sales", money(revenue)),
        ("Total profit", money(profit)),
        ("Overall profit margin", percent(margin)),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        set_font(cells[0].paragraphs[0].add_run(label))
        set_font(cells[1].paragraphs[0].add_run(value))
    set_table_geometry(table, [4500, 4860])
    add_caption(doc, "Table 1. Dataset profile and baseline financial results. Source: APL_Logistics.csv.")
    add_text(doc, "The dashboard reports total revenue, total profit, profit margin, and average discount as headline indicators. It provides market, customer, product, category, and discount analyses, including a product selector, top and bottom customer views, a category-market heatmap, and a discount scenario calculator. Because the source file does not contain an order-date field, the study does not make time-series claims or use trend forecasting.")

    add_heading(doc, "3. Exploratory Data Analysis and Findings")
    add_heading(doc, "3.1 Market performance", level=2)
    doc.add_picture(str(market_path), width=Inches(6.3))
    add_caption(doc, "Figure 1. Profit by market. Source: APL_Logistics.csv; author calculations.")
    top_market = market.sort_values(PROFIT, ascending=False).index[0]
    add_text(doc, f"Europe generated the highest total profit at {money(market.loc['Europe', PROFIT])}, followed by LATAM at {money(market.loc['LATAM', PROFIT])}. Together, these two markets account for a substantial share of overall sales and profit. This concentration creates an opportunity to protect successful commercial patterns in the leading regions while examining whether similar practices can be adapted for smaller markets.")

    add_heading(doc, "3.2 Customer value by segment", level=2)
    doc.add_picture(str(segment_path), width=Inches(6.3))
    add_caption(doc, "Figure 2. Profit by customer segment. Source: APL_Logistics.csv; author calculations.")
    highest_segment = segments[PROFIT].idxmax()
    lowest_segment = segments[PROFIT].idxmin()
    add_text(doc, f"Customer segments differ in their aggregate value. {highest_segment} generated the highest segment-level profit, while {lowest_segment} generated the lowest. The dashboard complements this aggregate view with top and bottom customer charts and an attention table for customers with above-median revenue but profit margins below 10%. These views support targeted pricing, retention, and account-management decisions.")

    add_heading(doc, "3.3 Discount and margin relationship", level=2)
    doc.add_picture(str(discount_path), width=Inches(6.3))
    add_caption(doc, "Figure 3. Profit margin by discount band. Source: APL_Logistics.csv; author calculations.")
    highest_band = discount.loc[discount['Margin %'].idxmax(), 'Discount band']
    lowest_band = discount.loc[discount['Margin %'].idxmin(), 'Discount band']
    add_text(doc, f"Discounting is a material commercial lever: the average order-item discount rate was {percent(df['Discount %'].mean())}. The margin comparison indicates that the {highest_band} band has the strongest aggregate margin, whereas the {lowest_band} band has the weakest. The dashboard therefore exposes discount rate and profitability together so commercial teams can check whether discount approvals are producing incremental value rather than simply reducing margin.")

    add_heading(doc, "3.4 Product category performance", level=2)
    top_categories = categories.head(5)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ["Category", "Sales", "Profit"]):
        set_cell_shading(cell, "F2F4F7")
        set_font(cell.paragraphs[0].add_run(value), bold=True, color=DARK_BLUE)
    for category, values in top_categories.iterrows():
        cells = table.add_row().cells
        set_font(cells[0].paragraphs[0].add_run(str(category)))
        set_font(cells[1].paragraphs[0].add_run(money(values['Sales'])))
        set_font(cells[2].paragraphs[0].add_run(money(values[PROFIT])))
    set_table_geometry(table, [4200, 2580, 2580])
    add_caption(doc, "Table 2. Five highest-profit product categories. Source: APL_Logistics.csv; author calculations.")
    add_text(doc, f"Fishing is the largest profit contributor ({money(top_categories.iloc[0][PROFIT])}), followed by Cleats and Camping & Hiking. These categories should be treated as strategic performance segments: they need disciplined availability, pricing, and promotion monitoring. Lower-volume or lower-margin categories should be reviewed using the dashboard before additional promotional spending is approved.")

    add_heading(doc, "4. Recommendations")
    add_text(doc, "Recommendation 1: Establish discount guardrails. Define category- and customer-segment-specific discount thresholds, require approval for high-discount orders, and monitor the resulting margin weekly. The objective is to preserve profitable demand while limiting margin dilution.", "Recommendation 1: ")
    add_text(doc, "Recommendation 2: Use customer-value tiers. Combine total profit, profit margin, and line-item volume to identify high-value accounts to protect and low-value accounts that require revised pricing, service, or account-management strategies.", "Recommendation 2: ")
    add_text(doc, "Recommendation 3: Protect high-value categories and markets. Prioritize availability and promotional discipline for Fishing, Cleats, Camping & Hiking, Europe, and LATAM. Pair this with category-level profit monitoring to prevent revenue growth that does not translate into margin.", "Recommendation 3: ")
    add_text(doc, "Recommendation 4: Institutionalize dashboard governance. Publish a weekly management view of revenue, profit, margin, discount, customer value, and category-market performance. Assign owners to investigate adverse movements and document decisions in an operational review cadence.", "Recommendation 4: ")

    add_heading(doc, "5. Executive Summary for Government Stakeholders")
    add_text(doc, "Reliable logistics networks support trade, employment, consumer access, and regional competitiveness. This project demonstrates how commercial data can be converted into transparent performance indicators for policy and management conversations. The dataset shows a commercially significant logistics operation with more than 180,000 order lines and sales exceeding $36 million. Its profitability, customer, product, and market views demonstrate how decision makers can move from revenue-focused reporting to margin-aware oversight.")
    add_text(doc, "For public-sector stakeholders, the dashboard is relevant as a model for evidence-based oversight. Aggregated views can reveal where infrastructure constraints, service-level gaps, or market disparities warrant additional investigation. The analysis should be used alongside qualitative carrier, warehouse, and customer evidence; it does not identify individual causes of delay or justify conclusions about any region without further validation.")

    add_heading(doc, "6. Limitations and Future Work")
    add_text(doc, "The study is descriptive and based on a single supplied dataset. The absence of order dates prevents seasonal, longitudinal, and forecasting analysis. Future work should add time stamps, fulfilment cost, customer satisfaction, returns, and campaign information. Those fields would support forecasting, customer-lifetime analysis, promotion measurement, and more robust causal or predictive modelling.")

    add_heading(doc, "7. Conclusion")
    add_text(doc, f"The APL Logistics dashboard converts {len(df):,} order-line records into a practical decision-support tool. EDA shows {money(revenue)} in sales, {money(profit)} in profit, a {percent(margin)} overall margin, strong value concentration in Europe and LATAM, and meaningful variation across customer segments. The most immediate actions are to strengthen discount governance, use customer-value tiers, protect high-value categories, and embed the dashboard in a recurring performance-review process.")

    add_heading(doc, "References")
    for reference in [
        "APL Logistics. (2026). APL_Logistics.csv [Data set provided for the Supply Chain Analytics Project].",
        "McKinney, W. (2024). pandas: powerful Python data analysis toolkit. https://pandas.pydata.org/",
        "Streamlit. (2026). Streamlit documentation. https://docs.streamlit.io/",
        "Plotly. (2026). Plotly Python graphing library documentation. https://plotly.com/python/",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        set_font(p.add_run(reference), size=10)

    doc.core_properties.title = "APL Logistics Profitability Intelligence"
    doc.core_properties.subject = "Supply chain analytics research paper"
    doc.core_properties.author = "[Student Name]"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
